import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class DocxGenerator:
    @staticmethod
    def generate_docx_bytes(title: str, text_content: str) -> io.BytesIO:
        doc = Document()
        
        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
            
        # Add Title
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Format text lines
        for paragraph_text in text_content.split("\n"):
            line = paragraph_text.strip()
            if line:
                if line.startswith("#"):
                    doc.add_heading(line.replace("#", "").strip(), level=2)
                elif line.startswith("- ") or line.startswith("* "):
                    doc.add_paragraph(line[2:].strip(), style='List Bullet')
                else:
                    p = doc.add_paragraph(line)
                    p.style.font.name = 'Calibri'
                    p.style.font.size = Pt(11)
                    
        # Save to byte buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
